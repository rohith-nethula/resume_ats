import streamlit as st
import os
from groq import Groq
from docx import Document
from docx.shared import Pt
import re
import json
from io import BytesIO

# Page config
st.set_page_config(
    page_title="ATS Resume Optimizer",
    page_icon="📄",
    layout="wide"
)

st.title("📄 ATS Resume Optimizer")
st.markdown("Optimize your resume for better ATS scores with AI (Powered by Groq - Free!)")

# Sidebar
with st.sidebar:
    st.header("⚙️ Settings")
    api_key = st.text_input("Groq API Key", type="password", help="Get free key from console.groq.com")
    
    if not api_key:
        st.info("👉 Get your FREE Groq API key from [console.groq.com](https://console.groq.com)\n\nNo credit card needed!")
    
    st.markdown("---")
    st.markdown("""
    **How it works:**
    1. Get free Groq API key
    2. Add key to sidebar
    3. Upload resume or paste text
    4. Paste job description
    5. Click Analyze
    6. Download optimized resume
    
    **Why Groq?**
    - ✅ Completely FREE
    - ✅ 10x faster than others
    - ✅ No credit card needed
    - ✅ Unlimited usage
    """)

# Check API key
if not api_key:
    st.warning("Please enter your Groq API key to get started")
    st.stop()

# Initialize Groq client
client = Groq(api_key=api_key)

# Session state for results
if 'analysis_complete' not in st.session_state:
    st.session_state.analysis_complete = False
    st.session_state.results = {}

def extract_text_from_docx(file):
    """Extract text from Word document"""
    try:
        doc = Document(file)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        return '\n'.join(text)
    except Exception as e:
        st.error(f"Error reading document: {e}")
        return None

def extract_json_from_response(response_text):
    """Extract JSON from Groq response"""
    try:
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            return json.loads(json_match.group())
    except:
        pass
    return None

def analyze_resume(resume_text, job_text):
    """Analyze resume against job description using Groq"""
    
    # Step 1: Analyze job
    with st.spinner("📊 Analyzing job description..."):
        job_response = client.chat.completions.create(
            model="llama-3.1-70b-versatile",  # Updated Groq model
            messages=[{
                "role": "user",
                "content": f"""Analyze this job description and extract:
1. Required technical skills
2. Required experience level
3. Key tools and technologies
4. Important keywords for ATS

Job Description:
{job_text}

Provide structured analysis."""
            }],
            temperature=0.3,
            max_tokens=1500
        )
        job_analysis = job_response.choices[0].message.content
    
    # Step 2: Compare resume to job
    with st.spinner("🔄 Comparing resume to job..."):
        comparison_response = client.chat.completions.create(
            model="llama-3.1-70b-versatile",
            messages=[
                {
                    "role": "user",
                    "content": f"Analyze this job:\n{job_text}"
                },
                {
                    "role": "assistant",
                    "content": job_analysis
                },
                {
                    "role": "user",
                    "content": f"""Compare this resume to the job:

Resume:
{resume_text}

Analyze:
1. Which required keywords are present?
2. Which are missing?
3. Estimated current ATS score (0-100)?
4. Where are the gaps?"""
                }
            ],
            temperature=0.3,
            max_tokens=1500
        )
        comparison = comparison_response.choices[0].message.content
    
    # Step 3: Generate suggestions
    with st.spinner("✨ Generating optimization suggestions..."):
        suggestions_response = client.chat.completions.create(
            model="llama-3.1-70b-versatile",
            messages=[
                {
                    "role": "user",
                    "content": f"Analyze this job:\n{job_text}"
                },
                {
                    "role": "assistant",
                    "content": job_analysis
                },
                {
                    "role": "user",
                    "content": f"""Compare this resume:\n{resume_text}"""
                },
                {
                    "role": "assistant",
                    "content": comparison
                },
                {
                    "role": "user",
                    "content": """Provide specific optimization suggestions as JSON format:
{"changes": [{"original": "exact text from resume", "improved": "better version with keywords", "section": "section name", "reason": "why this helps"}]}

Also estimate the new ATS score if these changes are implemented.

IMPORTANT: Provide the JSON first, then any additional explanation."""
                }
            ],
            temperature=0.3,
            max_tokens=2000
        )
        suggestions = suggestions_response.choices[0].message.content
    
    return {
        'job_analysis': job_analysis,
        'comparison': comparison,
        'suggestions': suggestions,
        'changes': extract_json_from_response(suggestions)
    }

# Main app
col1, col2 = st.columns(2)

with col1:
    st.subheader("Your Resume")
    resume_input = st.radio("Input method:", ["📝 Paste Text", "📄 Upload Word (.docx)"])
    
    resume_text = ""
    if resume_input == "📄 Upload Word (.docx)":
        resume_file = st.file_uploader("Upload resume", type=['docx'])
        if resume_file:
            resume_text = extract_text_from_docx(resume_file)
            if resume_text:
                st.success("✅ Resume loaded")
    else:
        resume_text = st.text_area("Paste resume:", height=250, placeholder="JOHN DOE\nSoftware Engineer\n\nEXPERIENCE\n...")

with col2:
    st.subheader("Job Description")
    job_text = st.text_area("Paste job description:", height=250, placeholder="Senior Software Engineer\n\nRequired:\n- 5+ years experience\n- Python and JavaScript\n...")

# Analyze button
if st.button("🔍 Analyze Resume", use_container_width=True, type="primary"):
    if not resume_text.strip():
        st.error("Please provide your resume")
    elif not job_text.strip():
        st.error("Please provide job description")
    else:
        try:
            results = analyze_resume(resume_text, job_text)
            st.session_state.results = results
            st.session_state.analysis_complete = True
            st.session_state.resume_text = resume_text
            st.success("✅ Analysis complete!")
        except Exception as e:
            st.error(f"Error during analysis: {str(e)}")

# Display results
if st.session_state.analysis_complete:
    st.markdown("---")
    
    tab1, tab2, tab3 = st.tabs(["📋 Analysis", "💡 Suggestions", "📥 Download"])
    
    with tab1:
        with st.expander("Job Requirements", expanded=True):
            st.markdown(st.session_state.results['job_analysis'])
        
        with st.expander("Resume Comparison", expanded=True):
            st.markdown(st.session_state.results['comparison'])
    
    with tab2:
        suggestions_text = st.session_state.results['suggestions']
        st.markdown(suggestions_text)
        
        # Show structured changes
        changes = st.session_state.results.get('changes')
        if changes and 'changes' in changes:
            st.subheader("Specific Changes:")
            for i, change in enumerate(changes['changes'], 1):
                with st.container():
                    st.write(f"**Change {i}** ({change.get('section', 'General')})")
                    col1, col2 = st.columns(2)
                    with col1:
                        st.write(f"**Original:** {change.get('original', '')}")
                    with col2:
                        st.write(f"**Improved:** {change.get('improved', '')}")
                    st.caption(f"Why: {change.get('reason', '')}")
                    st.divider()
    
    with tab3:
        # Create optimized resume
        resume_text = st.session_state.get('resume_text', '')
        changes = st.session_state.results.get('changes')
        
        if resume_text:
            # Apply changes
            modified_text = resume_text
            if changes and 'changes' in changes:
                for change in changes['changes']:
                    original = change.get('original', '')
                    improved = change.get('improved', '')
                    if original and improved:
                        modified_text = modified_text.replace(original, improved)
            
            # Download options
            col1, col2 = st.columns(2)
            
            with col1:
                # Create Word document
                try:
                    doc = Document()
                    title = doc.add_paragraph()
                    title_run = title.add_run("OPTIMIZED RESUME")
                    title_run.bold = True
                    title_run.font.size = Pt(14)
                    
                    for line in modified_text.split('\n'):
                        if line.strip():
                            doc.add_paragraph(line)
                    
                    doc_bytes = BytesIO()
                    doc.save(doc_bytes)
                    doc_bytes.seek(0)
                    
                    st.download_button(
                        "📄 Download as Word (.docx)",
                        doc_bytes.getvalue(),
                        "optimized_resume.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                except Exception as e:
                    st.error(f"Error creating document: {e}")
            
            with col2:
                st.download_button(
                    "📋 Download as Text (.txt)",
                    modified_text,
                    "optimized_resume.txt",
                    "text/plain",
                    use_container_width=True
                )
            
            # Show text
            st.text_area("Optimized Resume:", modified_text, height=300, disabled=True)

# Footer
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: gray; font-size: 12px;'>
<p>ATS Resume Optimizer | Powered by Groq AI | FREE forever!</p>
<p>Get API key: <a href='https://console.groq.com'>console.groq.com</a> (No credit card needed)</p>
</div>
""", unsafe_allow_html=True)
